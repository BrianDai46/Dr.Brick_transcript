import docx
import re
import pandas as pd


def doc_extract(document):
    runs_list = []
    doc = docx.Document(document)
    for i in doc.paragraphs:
        for run in i.runs:
            runs_list.append(run.text)
    return runs_list


def doc_extract_encoded(document):
    text_holder = []
    doc = docx.Document(document)
    paragraphs_count = 0
    runs_count = 0
    index_num = 0
    for i in doc.paragraphs:
        paragraphs_count += 1
        for run in i.runs:
            text_holder.append(run.text)
            text_holder[index_num] = [text_holder[index_num], paragraphs_count, runs_count]
            runs_count += 1
            index_num += 1
        runs_count = 0
    return text_holder


def search_name(runs_list):
    name_pattern = re.compile("^[\w '0-9]{3,25}\:")
    name_list = []
    for i in range(len(runs_list)):
        if name_pattern.match(runs_list[i][0]) is not None:
            name_list.append(re.sub(" \[.*", "", runs_list[i][0]))
    return name_list


def search_time(runs_list):
    time_list = []
    time_pattern = re.compile("([0-5]\d):([0-5]\d):([0-5]0)\]")
    time_list_temp = list(filter(time_pattern.search, runs_list))
    for i in range(len(time_list_temp)):
        time_list.append(re.sub(".+\[", "[", time_list_temp[i]))
    return time_list


def search_text(runs_list):
    text_pattern = re.compile("^[\w0-9 .,?']+$")
    text_list = list(filter(text_pattern.search, runs_list))
    return text_list


# def time_loca_track(time_list, encoded_list):


def text_talk_turn_track(text_list, encoded_list):
    talk_turn = []
    for i in range(len(encoded_list)):
        for j in range(len(text_list)):
            if encoded_list[i][0] != text_list[j]:
                continue
            else:
                talk_turn.append(encoded_list[i][1])
    return talk_turn


def text_segment_track(talk_turn):
    segment_list = []
    seg_count = 1
    for i in range(len(talk_turn)):
        last_turn = talk_turn[i - 1]
        if talk_turn[i] == last_turn:
            seg_count += 1
        else:
            seg_count = 1
        segment_list.append(seg_count)
    return segment_list


def text_speaker_track(talk_turn, name_list):
    speaker_list = []
    for i in range(len(talk_turn)):
        speaker_list.append(name_list[talk_turn[i] - 1])
    return speaker_list


# def get_time_posi(encoded_list,)
# print(text_segment_track(text_talk_turn_track(search_text(doc_extract('Example transcript Short.docx'))))
print(doc_extract_encoded('Example transcript Short.docx'))
print(search_time(doc_extract('Example transcript Short.docx')))

column_names = ["Time", "Talk_turn", "Segment", "Speaker", "Text"]
df = pd.DataFrame(columns=column_names)
df.Text = search_text(doc_extract('Example transcript Short.docx'))
df.Talk_turn = text_talk_turn_track(search_text(doc_extract('Example transcript Short.docx')),
                                    doc_extract_encoded('Example transcript Short.docx'))
df.Speaker = text_speaker_track(text_talk_turn_track(search_text(doc_extract('Example transcript Short.docx')),
                                                     doc_extract_encoded('Example transcript Short.docx')),
                                search_name(doc_extract_encoded('Example transcript Short.docx')))
df.Segment = text_segment_track(text_talk_turn_track(search_text(doc_extract('Example transcript Short.docx')),
                                                     doc_extract_encoded('Example transcript Short.docx')))
df.to_csv('out.csv')
